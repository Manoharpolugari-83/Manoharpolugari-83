# ============================================================
# AI RESUME SCREENER & OPTIMIZER - COMPLETE ENHANCED VERSION
# Using Groq API with LLaMA 3.3-70b
# ============================================================
# SETUP:
# 1. pip install flask flask-cors groq PyPDF2 python-docx
# 2. Get FREE API key: https://console.groq.com/keys
# 3. Set your key below (GROQ_API_KEY = "your_key")
# 4. python resume.py
# 5. Open http://localhost:5000
# ============================================================

import os, json, re, io, traceback
from datetime import datetime
from flask import Flask, request, jsonify, Response
from flask_cors import CORS
from werkzeug.utils import secure_filename

try:
    from groq import Groq
except ImportError:
    os.system("pip install groq")
    from groq import Groq

try:
    import PyPDF2
except ImportError:
    os.system("pip install PyPDF2")
    import PyPDF2

try:
    import docx
except ImportError:
    os.system("pip install python-docx")
    import docx

# ============================================================
# ✅ SET YOUR GROQ API KEY HERE
# ============================================================
GROQ_API_KEY = "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"

FALLBACK_MODELS = [
    "llama-3.3-70b-versatile",
    "llama-3.1-8b-instant",
    "llama3-70b-8192",
    "llama3-8b-8192",
    "gemma2-9b-it",
]
MODEL_NAME = FALLBACK_MODELS[0]
# ============================================================

app = Flask(__name__)
app.secret_key = "resume_ai_2025_secret"
CORS(app)
os.makedirs("uploads", exist_ok=True)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

client = None
if GROQ_API_KEY and GROQ_API_KEY != "your_groq_api_key_here":
    client = Groq(api_key=GROQ_API_KEY)
    print("✅ Groq API Ready")
else:
    print("❌ Set GROQ_API_KEY first!")

# ── Helpers ──────────────────────────────────────────────────
def allowed_file(f): return "." in f and f.rsplit(".",1)[1].lower() in {"pdf","docx","txt"}

def read_pdf(stream):
    r = PyPDF2.PdfReader(stream); t=""
    for p in r.pages:
        x=p.extract_text()
        if x: t+=x+"\n"
    return t.strip()

def read_docx(stream):
    d=docx.Document(stream); t=""
    for p in d.paragraphs: t+=p.text+"\n"
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells: t+=cell.text+" "
        t+="\n"
    return t.strip()

def read_file(file):
    ext=secure_filename(file.filename).rsplit(".",1)[1].lower()
    buf=io.BytesIO(file.read()); file.seek(0)
    if ext=="pdf": return read_pdf(buf)
    elif ext=="docx": return read_docx(buf)
    elif ext=="txt": return buf.read().decode("utf-8","ignore").strip()
    raise ValueError(f"Unsupported: {ext}")

def call_ai(prompt, retries=2):
    if not client: raise ValueError("Groq API key not configured!")
    last_err=None
    for model in FALLBACK_MODELS:
        for attempt in range(retries):
            try:
                print(f"🤖 {model} attempt {attempt+1}")
                r=client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role":"system","content":"You are an expert resume analyst. Return ONLY valid JSON. No markdown, no ```json, no extra text. Just raw JSON."},
                        {"role":"user","content":prompt}
                    ],
                    temperature=0.4, max_tokens=8000
                )
                print(f"✅ Success: {model}")
                return r.choices[0].message.content
            except Exception as e:
                last_err=e; s=str(e).lower()
                print(f"⚠ {model}: {str(e)[:80]}")
                if any(k in s for k in ["decommission","not found","does not exist","deprecated"]): break
    raise ValueError(f"All models failed. Last: {last_err}")

def parse_json(text):
    if not text: return {}
    text=re.sub(r"```(?:json)?\s*","",text)
    text=re.sub(r"```\s*$","",text).strip()
    try: return json.loads(text)
    except: pass
    try:
        s=text.find("{"); e=text.rfind("}")+1
        if s>=0 and e>s: return json.loads(text[s:e])
    except: pass
    return {"raw_response":text}

# ── HTML Template ─────────────────────────────────────────────
HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>AI Resume Screener Pro</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap" rel="stylesheet">
<link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.1/css/all.min.css">
<style>
/* ── Reset & Variables ── */
*{margin:0;padding:0;box-sizing:border-box}
:root{
  --p:#7c3aed;--pl:#a78bfa;--p2:#6d28d9;
  --s:#0ea5e9;--ok:#10b981;--warn:#f59e0b;--err:#ef4444;
  --pink:#ec4899;--teal:#14b8a6;
  --bg:#060918;--bg2:#0d1117;--card:#111827;--card2:#1a2332;--cardh:#1f2d42;
  --t1:#f8fafc;--t2:#94a3b8;--t3:#475569;--bdr:#1e293b;
  --g1:linear-gradient(135deg,#7c3aed,#a855f7,#ec4899);
  --g2:linear-gradient(135deg,#0ea5e9,#7c3aed);
  --g3:linear-gradient(135deg,#10b981,#0ea5e9);
  --g4:linear-gradient(135deg,#f59e0b,#ef4444);
  --gcard:linear-gradient(145deg,#111827,#1a2332);
  --shadow:0 4px 24px rgba(0,0,0,.4);
  --shadow2:0 8px 40px rgba(124,58,237,.15);
}

body{font-family:'Inter',sans-serif;background:var(--bg);color:var(--t1);min-height:100vh;overflow-x:hidden}

/* ── Animated BG ── */
.bg-mesh{position:fixed;inset:0;z-index:0;overflow:hidden;pointer-events:none}
.bg-mesh::before{
  content:'';position:absolute;width:800px;height:800px;
  background:radial-gradient(circle,rgba(124,58,237,.08) 0%,transparent 70%);
  top:-200px;left:-200px;animation:pulse 8s ease-in-out infinite
}
.bg-mesh::after{
  content:'';position:absolute;width:600px;height:600px;
  background:radial-gradient(circle,rgba(14,165,233,.06) 0%,transparent 70%);
  bottom:-100px;right:-100px;animation:pulse 8s ease-in-out infinite reverse
}
.bg-orb{
  position:fixed;width:300px;height:300px;border-radius:50%;
  background:radial-gradient(circle,rgba(236,72,153,.05) 0%,transparent 70%);
  top:40%;left:50%;transform:translate(-50%,-50%);
  animation:orb 12s ease-in-out infinite;pointer-events:none;z-index:0
}
@keyframes pulse{0%,100%{transform:scale(1) translate(0,0)}50%{transform:scale(1.1) translate(30px,20px)}}
@keyframes orb{0%,100%{transform:translate(-50%,-50%) scale(1)}50%{transform:translate(-30%,-60%) scale(1.3)}}

/* ── Layout ── */
.wrap{position:relative;z-index:1;max-width:1500px;margin:0 auto;padding:24px 20px}

/* ── Header ── */
.hdr{text-align:center;padding:50px 20px 35px;position:relative}
.hdr-logo{
  display:inline-flex;align-items:center;gap:14px;
  margin-bottom:20px
}
.hdr-icon{
  width:60px;height:60px;border-radius:16px;
  background:var(--g1);display:flex;align-items:center;
  justify-content:center;font-size:1.6rem;
  box-shadow:0 8px 24px rgba(124,58,237,.4)
}
.hdr h1{
  font-size:3.2rem;font-weight:900;letter-spacing:-2px;
  background:var(--g1);-webkit-background-clip:text;
  -webkit-text-fill-color:transparent;background-clip:text;
  line-height:1.1
}
.hdr-sub{color:var(--t2);font-size:1.05rem;margin-top:10px;font-weight:400;letter-spacing:.3px}
.hdr-badges{display:flex;gap:10px;justify-content:center;flex-wrap:wrap;margin-top:18px}
.bdg{
  display:inline-flex;align-items:center;gap:7px;
  padding:7px 16px;border-radius:50px;font-size:.82rem;font-weight:600;
  border:1px solid;letter-spacing:.3px;
}
.bdg-purple{background:rgba(124,58,237,.15);border-color:rgba(124,58,237,.4);color:var(--pl)}
.bdg-green{background:rgba(16,185,129,.1);border-color:rgba(16,185,129,.3);color:#34d399}
.bdg-blue{background:rgba(14,165,233,.1);border-color:rgba(14,165,233,.3);color:#38bdf8}

/* ── Tabs ── */
.tabs{
  display:flex;gap:6px;justify-content:center;
  margin-bottom:28px;flex-wrap:wrap;
  background:var(--card);border:1px solid var(--bdr);
  border-radius:16px;padding:8px;max-width:700px;margin-left:auto;margin-right:auto;
  margin-bottom:32px
}
.tb{
  flex:1;min-width:120px;padding:12px 20px;
  border:none;background:transparent;color:var(--t2);
  border-radius:10px;cursor:pointer;font-size:.88rem;font-weight:600;
  transition:all .25s;display:flex;align-items:center;
  justify-content:center;gap:8px;font-family:'Inter',sans-serif;
  white-space:nowrap;letter-spacing:.3px
}
.tb:hover{color:var(--t1);background:var(--cardh)}
.tb.on{
  background:var(--g1);color:#fff;
  box-shadow:0 4px 16px rgba(124,58,237,.35)
}

/* ── Tab content ── */
.tc{display:none;animation:fadeup .35s ease}
.tc.on{display:block}
@keyframes fadeup{from{opacity:0;transform:translateY(12px)}to{opacity:1;transform:translateY(0)}}

/* ── Cards ── */
.cd{
  background:var(--gcard);border:1px solid var(--bdr);
  border-radius:20px;padding:28px;margin-bottom:20px;
  box-shadow:var(--shadow);transition:border-color .3s
}
.cd:hover{border-color:rgba(124,58,237,.25)}
.cd-glass{
  background:rgba(17,24,39,.6);backdrop-filter:blur(12px);
  border:1px solid rgba(255,255,255,.06)
}
.ct{
  font-size:1.15rem;font-weight:700;margin-bottom:22px;
  display:flex;align-items:center;gap:10px;letter-spacing:-.3px
}
.ct .icon{
  width:36px;height:36px;border-radius:10px;
  background:rgba(124,58,237,.2);display:flex;
  align-items:center;justify-content:center;
  color:var(--pl);font-size:.95rem;flex-shrink:0
}

/* ── Form Elements ── */
.fg{margin-bottom:18px}
.fg label{display:block;margin-bottom:8px;font-weight:600;font-size:.88rem;color:var(--t2);letter-spacing:.3px;text-transform:uppercase}
.fg label .rq{color:var(--err);margin-left:2px}
textarea,input[type=text],input[type=email],input[type=tel],select{
  width:100%;padding:14px 16px;
  background:rgba(6,9,24,.6);border:1.5px solid var(--bdr);
  border-radius:12px;color:var(--t1);font-family:'Inter',sans-serif;
  font-size:.93rem;transition:all .25s;resize:vertical;
  line-height:1.6
}
textarea:focus,input:focus,select:focus{
  outline:none;border-color:var(--p);
  box-shadow:0 0 0 4px rgba(124,58,237,.12);
  background:rgba(6,9,24,.8)
}
textarea{min-height:140px}
textarea.lg{min-height:260px}
select option{background:var(--card)}

/* ── File Upload ── */
.fu{
  border:2px dashed var(--bdr);border-radius:16px;
  padding:36px 20px;text-align:center;cursor:pointer;
  transition:all .3s;position:relative;overflow:hidden;
  background:rgba(6,9,24,.3)
}
.fu:hover,.fu.dg{
  border-color:var(--p);
  background:rgba(124,58,237,.05);
}
.fu-icon{
  width:56px;height:56px;border-radius:14px;
  background:rgba(124,58,237,.15);display:flex;
  align-items:center;justify-content:center;
  margin:0 auto 14px;font-size:1.4rem;color:var(--pl);
  transition:all .3s
}
.fu:hover .fu-icon{background:rgba(124,58,237,.25);transform:scale(1.05)}
.fu p{color:var(--t2);font-size:.9rem;font-weight:500}
.fu .sm{font-size:.78rem;color:var(--t3);margin-top:5px}
.fu .fn{
  margin-top:12px;color:var(--ok);font-weight:600;
  font-size:.88rem;display:none;
  padding:8px 16px;background:rgba(16,185,129,.1);
  border-radius:8px;border:1px solid rgba(16,185,129,.2)
}
.fu .fn.show{display:inline-block}
.fu input[type=file]{position:absolute;inset:0;opacity:0;cursor:pointer}

/* ── Buttons ── */
.btn{
  padding:14px 30px;border:none;border-radius:12px;
  font-size:.95rem;font-weight:700;cursor:pointer;
  transition:all .25s;display:inline-flex;
  align-items:center;gap:9px;font-family:'Inter',sans-serif;
  letter-spacing:.3px;position:relative;overflow:hidden
}
.btn::after{
  content:'';position:absolute;inset:0;
  background:rgba(255,255,255,0);transition:background .2s
}
.btn:hover::after{background:rgba(255,255,255,.08)}
.btn:active{transform:scale(.98)}
.btn:disabled{opacity:.45;cursor:not-allowed}
.btn:disabled:hover::after{background:none}
.bp{background:var(--g1);color:#fff;box-shadow:0 4px 20px rgba(124,58,237,.35)}
.bp:hover{box-shadow:0 6px 28px rgba(124,58,237,.5);transform:translateY(-1px)}
.bg2{background:var(--g3);color:#fff;box-shadow:0 4px 20px rgba(16,185,129,.3)}
.bg2:hover{box-shadow:0 6px 28px rgba(16,185,129,.45);transform:translateY(-1px)}
.bs{
  background:var(--cardh);color:var(--t1);
  border:1.5px solid var(--bdr)
}
.bs:hover{border-color:var(--p);transform:translateY(-1px)}
.btnr{display:flex;gap:10px;flex-wrap:wrap}
.btn-ctr{text-align:center;margin:24px 0}

/* ── Grid ── */
.g2{display:grid;grid-template-columns:1fr 1fr;gap:20px}
.g3{display:grid;grid-template-columns:repeat(3,1fr);gap:16px}
.g4{display:grid;grid-template-columns:repeat(4,1fr);gap:14px}
.ir{display:grid;grid-template-columns:1fr 1fr;gap:16px}
@media(max-width:950px){.g2{grid-template-columns:1fr}}
@media(max-width:768px){
  .g2,.g3,.g4,.ir{grid-template-columns:1fr}
  .hdr h1{font-size:2.2rem}
  .tabs{flex-direction:column;max-width:100%}
  .tb{flex:none;min-width:unset}
  .wrap{padding:16px}
}

/* ── Result Section ── */
.rs{display:none;animation:fadeup .5s ease}
.rs.sh{display:block}

/* ── Score Hero ── */
.score-hero{
  text-align:center;padding:40px 20px 32px;
  background:var(--gcard);border:1px solid var(--bdr);
  border-radius:20px;margin-bottom:20px;position:relative;overflow:hidden
}
.score-hero::before{
  content:'';position:absolute;inset:0;
  background:radial-gradient(circle at 50% 0%,rgba(124,58,237,.08) 0%,transparent 60%);
  pointer-events:none
}
.score-ring-wrap{
  display:inline-flex;align-items:center;
  justify-content:center;margin-bottom:20px;
  position:relative
}
.score-ring{
  width:160px;height:160px;position:relative;
  display:flex;align-items:center;justify-content:center
}
.score-ring svg{position:absolute;inset:0;transform:rotate(-90deg)}
.score-ring-inner{
  position:relative;z-index:1;text-align:center
}
.score-val{font-size:3rem;font-weight:900;line-height:1}
.score-lbl{font-size:.78rem;color:var(--t2);font-weight:600;letter-spacing:.5px;text-transform:uppercase;margin-top:3px}
.fit-chip{
  display:inline-flex;align-items:center;gap:8px;
  padding:10px 24px;border-radius:50px;
  font-size:1rem;font-weight:800;text-transform:uppercase;
  letter-spacing:1.5px;margin-bottom:14px
}
.fit-excellent{background:rgba(16,185,129,.15);color:#34d399;border:2px solid rgba(16,185,129,.4)}
.fit-good{background:rgba(14,165,233,.15);color:#38bdf8;border:2px solid rgba(14,165,233,.4)}
.fit-moderate{background:rgba(245,158,11,.15);color:#fbbf24;border:2px solid rgba(245,158,11,.4)}
.fit-low,.fit-poor{background:rgba(239,68,68,.15);color:#f87171;border:2px solid rgba(239,68,68,.4)}
.hero-summary{
  color:var(--t2);max-width:680px;margin:0 auto;
  line-height:1.75;font-size:.95rem
}
.prob-badge{
  display:inline-flex;align-items:center;gap:6px;
  margin-top:14px;padding:7px 18px;border-radius:50px;
  font-size:.85rem;font-weight:700;letter-spacing:.5px
}
.prob-high{background:rgba(16,185,129,.15);color:#34d399;border:1px solid rgba(16,185,129,.3)}
.prob-medium{background:rgba(245,158,11,.15);color:#fbbf24;border:1px solid rgba(245,158,11,.3)}
.prob-low{background:rgba(239,68,68,.15);color:#f87171;border:1px solid rgba(239,68,68,.3)}

/* ── Stat Cards ── */
.stat-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-bottom:20px}
@media(max-width:768px){.stat-grid{grid-template-columns:repeat(2,1fr)}}
.stat-card{
  background:var(--card2);border:1px solid var(--bdr);
  border-radius:14px;padding:18px;text-align:center;
  transition:all .25s;position:relative;overflow:hidden
}
.stat-card::before{
  content:'';position:absolute;top:0;left:0;right:0;height:3px;
  border-radius:14px 14px 0 0
}
.stat-card.c-purple::before{background:var(--g1)}
.stat-card.c-blue::before{background:var(--g2)}
.stat-card.c-green::before{background:var(--g3)}
.stat-card.c-orange::before{background:var(--g4)}
.stat-card:hover{transform:translateY(-2px);border-color:rgba(124,58,237,.3)}
.stat-v{font-size:2.2rem;font-weight:900;margin-bottom:4px;line-height:1}
.stat-l{font-size:.75rem;color:var(--t3);font-weight:600;letter-spacing:.5px;text-transform:uppercase}
.stat-bar{height:4px;border-radius:2px;margin-top:10px;background:var(--bdr);overflow:hidden}
.stat-bar-fill{height:100%;border-radius:2px;transition:width 1.2s ease}

/* ── Progress Bars ── */
.sec-hdr{
  display:flex;align-items:center;gap:10px;
  margin-bottom:16px;padding-bottom:12px;
  border-bottom:1px solid var(--bdr)
}
.sec-hdr .ico{
  width:32px;height:32px;border-radius:8px;
  display:flex;align-items:center;justify-content:center;font-size:.85rem
}
.sec-hdr h3{font-size:1rem;font-weight:700;letter-spacing:-.2px}
.pi{margin-bottom:14px}
.ph{display:flex;justify-content:space-between;align-items:center;margin-bottom:6px}
.ph-label{font-size:.85rem;color:var(--t2);font-weight:500}
.ph-val{font-size:.85rem;font-weight:700;color:var(--t1)}
.pbar{height:8px;background:rgba(30,41,59,.6);border-radius:6px;overflow:hidden}
.pfill{
  height:100%;border-radius:6px;
  transition:width 1.3s cubic-bezier(.25,.46,.45,.94)
}
.pfill.grn{background:var(--g3)}
.pfill.blu{background:var(--g2)}
.pfill.ylw{background:linear-gradient(90deg,#f59e0b,#fbbf24)}
.pfill.rd{background:linear-gradient(90deg,#ef4444,#f87171)}
.pfill.pur{background:var(--g1)}

/* ── Tags ── */
.tags{display:flex;flex-wrap:wrap;gap:8px}
.tag{
  padding:6px 13px;border-radius:8px;
  font-size:.8rem;font-weight:600;letter-spacing:.2px;
  transition:all .2s
}
.tag:hover{transform:scale(1.03)}
.tg{background:rgba(16,185,129,.1);color:#34d399;border:1px solid rgba(16,185,129,.25)}
.tr{background:rgba(239,68,68,.1);color:#f87171;border:1px solid rgba(239,68,68,.25)}
.tb2{background:rgba(14,165,233,.1);color:#38bdf8;border:1px solid rgba(14,165,233,.25)}
.ty{background:rgba(245,158,11,.1);color:#fbbf24;border:1px solid rgba(245,158,11,.25)}
.tp{background:rgba(124,58,237,.1);color:var(--pl);border:1px solid rgba(124,58,237,.25)}
.tpk{background:rgba(236,72,153,.1);color:#f472b6;border:1px solid rgba(236,72,153,.25)}

/* ── Lists ── */
.rl{list-style:none;padding:0}
.rl li{
  padding:12px 16px;border-left:3px solid var(--p);
  margin-bottom:8px;background:rgba(6,9,24,.5);
  border-radius:0 10px 10px 0;font-size:.9rem;
  color:var(--t2);transition:all .2s;line-height:1.55;
  cursor:default
}
.rl li:hover{background:rgba(124,58,237,.08);color:var(--t1);border-left-color:var(--pl)}
.rl.grl li{border-left-color:var(--ok)}
.rl.grl li:hover{background:rgba(16,185,129,.06);border-left-color:#34d399}
.rl.rrl li{border-left-color:var(--err)}
.rl.rrl li:hover{background:rgba(239,68,68,.06);border-left-color:#f87171}
.rl.yrl li{border-left-color:var(--warn)}
.rl.yrl li:hover{background:rgba(245,158,11,.06);border-left-color:#fbbf24}

/* ── Correction Cards ── */
.cc{
  background:rgba(6,9,24,.5);border:1px solid var(--bdr);
  border-radius:14px;padding:18px;margin-bottom:14px;
  transition:all .2s
}
.cc:hover{border-color:rgba(124,58,237,.25);background:rgba(6,9,24,.7)}
.cc .og{
  color:var(--err);text-decoration:line-through;
  margin-bottom:8px;font-size:.9rem;line-height:1.55;
  padding:8px 12px;background:rgba(239,68,68,.06);
  border-radius:8px
}
.cc .cr{
  color:var(--ok);font-weight:500;margin-bottom:8px;
  font-size:.9rem;line-height:1.55;
  padding:8px 12px;background:rgba(16,185,129,.06);
  border-radius:8px
}
.cc .ex{color:var(--t3);font-size:.83rem;font-style:italic;line-height:1.5;margin-top:6px}
.prb{
  display:inline-flex;align-items:center;gap:4px;
  padding:3px 10px;border-radius:6px;
  font-size:.72rem;font-weight:700;text-transform:uppercase;letter-spacing:.5px
}
.prh{background:rgba(239,68,68,.15);color:#f87171;border:1px solid rgba(239,68,68,.25)}
.prm{background:rgba(245,158,11,.15);color:#fbbf24;border:1px solid rgba(245,158,11,.25)}
.prl{background:rgba(16,185,129,.15);color:#34d399;border:1px solid rgba(16,185,129,.25)}

/* ── Resume Preview ── */
.rpv{
  background:#fff;color:#1a1a1a;padding:48px;
  border-radius:16px;max-height:900px;overflow-y:auto;
  font-size:.93rem;line-height:1.7;
  box-shadow:0 4px 40px rgba(0,0,0,.4)
}

/* ── Loading ── */
.lo{
  display:none;position:fixed;inset:0;
  background:rgba(6,9,24,.92);backdrop-filter:blur(8px);
  z-index:2000;justify-content:center;
  align-items:center;flex-direction:column
}
.lo.sh{display:flex}
.lo-inner{text-align:center;padding:40px}
.ldr{
  width:64px;height:64px;border-radius:50%;
  border:3px solid var(--bdr);border-top:3px solid var(--p);
  border-right:3px solid var(--pink);
  animation:spin 1s linear infinite;margin:0 auto 20px
}
@keyframes spin{to{transform:rotate(360deg)}}
.lo-text{font-size:1.1rem;font-weight:600;margin-bottom:6px}
.lo-sub{color:var(--t3);font-size:.88rem}
.lo-dots{color:var(--pl);font-size:1.5rem;margin-top:8px;letter-spacing:4px;animation:dots 1.5s steps(4) infinite}
@keyframes dots{0%{content:''}25%{content:'.'}50%{content:'..'}75%{content:'...'}100%{content:''}}

/* ── Toast ── */
.toast{
  position:fixed;bottom:28px;right:28px;
  padding:14px 20px;border-radius:14px;
  color:#fff;font-weight:600;z-index:3000;
  display:none;box-shadow:0 8px 30px rgba(0,0,0,.4);
  font-size:.9rem;max-width:340px;line-height:1.4;
  border:1px solid rgba(255,255,255,.1);
  backdrop-filter:blur(10px)
}
.toast.sh{display:flex;align-items:center;gap:10px;animation:tostin .3s ease}
@keyframes tostin{from{transform:translateX(80px);opacity:0}to{transform:translateX(0);opacity:1}}
.tok{background:rgba(16,185,129,.9)}
.ter{background:rgba(239,68,68,.9)}
.twn{background:rgba(245,158,11,.9)}

/* ── Divider ── */
.odv{
  display:flex;align-items:center;gap:14px;
  margin:16px 0;color:var(--t3);font-size:.82rem;font-weight:600;
  text-transform:uppercase;letter-spacing:.5px
}
.odv::before,.odv::after{content:'';flex:1;height:1px;background:var(--bdr)}

/* ── Info Box ── */
.ibox{
  padding:14px 18px;border-radius:12px;margin-bottom:18px;
  font-size:.88rem;line-height:1.6;display:flex;align-items:flex-start;gap:12px
}
.ibox i{margin-top:2px;flex-shrink:0}
.ibox.inf{background:rgba(124,58,237,.08);border:1px solid rgba(124,58,237,.2);color:var(--pl)}
.ibox.ok{background:rgba(16,185,129,.08);border:1px solid rgba(16,185,129,.2);color:#34d399}
.ibox.warn{background:rgba(245,158,11,.08);border:1px solid rgba(245,158,11,.2);color:#fbbf24}
.ibox.err{background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.2);color:#f87171}

/* ── Comparison Card ── */
.cmp-card{
  background:var(--card2);border:1px solid var(--bdr);
  border-radius:16px;padding:22px;margin-bottom:14px;
  transition:all .25s
}
.cmp-card:hover{border-color:rgba(124,58,237,.3);transform:translateY(-1px)}
.cmp-rank{
  display:flex;align-items:center;justify-content:space-between;
  margin-bottom:14px
}
.cmp-title{font-size:1.05rem;font-weight:700}
.mini-score{
  width:58px;height:58px;border-radius:50%;
  display:flex;flex-direction:column;align-items:center;
  justify-content:center;font-size:1.2rem;font-weight:900
}

/* ── Error Display ── */
.err-box{
  padding:20px;background:rgba(239,68,68,.08);
  border:1px solid rgba(239,68,68,.25);border-radius:14px;
  color:#f87171;font-size:.9rem;line-height:1.7
}
.err-box a{color:#f87171;text-decoration:underline}

/* ── Scrollbar ── */
::-webkit-scrollbar{width:6px}
::-webkit-scrollbar-track{background:transparent}
::-webkit-scrollbar-thumb{background:var(--bdr);border-radius:3px}
::-webkit-scrollbar-thumb:hover{background:var(--t3)}

/* ── Responsive tweaks ── */
@media(max-width:600px){
  .score-ring{width:130px;height:130px}
  .score-val{font-size:2.4rem}
  .hdr h1{font-size:1.8rem}
  .btn{padding:12px 20px;font-size:.88rem}
  .cd{padding:20px}
}
</style>
</head>
<body>
<div class="bg-mesh"></div>
<div class="bg-orb"></div>

<!-- Loading -->
<div class="lo" id="lo">
  <div class="lo-inner">
    <div class="ldr"></div>
    <div class="lo-text" id="loT">Analyzing your resume...</div>
    <div class="lo-sub" id="loS">Powered by Groq · LLaMA 3.3 · 70B</div>
    <div class="lo-dots">···</div>
  </div>
</div>

<!-- Toast -->
<div class="toast" id="toast"></div>

<div class="wrap">

<!-- Header -->
<div class="hdr">
  <div class="hdr-logo">
    <div class="hdr-icon">🤖</div>
    <h1>AI Resume Screener</h1>
  </div>
  <p class="hdr-sub">Screen, score, correct &amp; generate ATS-optimized resumes using advanced AI</p>
  <div class="hdr-badges">
    <span class="bdg bdg-purple"><i class="fas fa-bolt"></i> Powered by Groq</span>
    <span class="bdg bdg-green" id="modelBdg"><i class="fas fa-microchip"></i> Connecting...</span>
    <span class="bdg bdg-blue"><i class="fas fa-shield-alt"></i> ATS Optimized</span>
  </div>
</div>

<!-- Tabs -->
<div class="tabs">
  <button class="tb on" onclick="stab('screen',this)"><i class="fas fa-search"></i> Screen &amp; Score</button>
  <button class="tb" onclick="stab('correct',this)"><i class="fas fa-spell-check"></i> Corrections</button>
  <button class="tb" onclick="stab('generate',this)"><i class="fas fa-magic"></i> Generate</button>
  <button class="tb" onclick="stab('compare',this)"><i class="fas fa-balance-scale"></i> Compare</button>
</div>

<!-- ══════════════ TAB 1: SCREEN ══════════════ -->
<div class="tc on" id="tab-screen">
  <div class="ibox inf">
    <i class="fas fa-info-circle"></i>
    <span>Upload or paste your resume, add the job description, and receive a comprehensive ATS compatibility score with keyword analysis, section scoring, and actionable recommendations.</span>
  </div>
  <div class="g2">
    <div>
      <div class="cd">
        <div class="ct"><div class="icon"><i class="fas fa-briefcase"></i></div> Job Description</div>
        <div class="fg">
          <label>Full job description <span class="rq">*</span></label>
          <textarea class="lg" id="sJD" placeholder="Paste the complete job description here. Include all requirements, responsibilities, and qualifications for the most accurate scoring..."></textarea>
        </div>
      </div>
    </div>
    <div>
      <div class="cd">
        <div class="ct"><div class="icon"><i class="fas fa-file-alt"></i></div> Your Resume</div>
        <div class="fu" id="sFU">
          <div class="fu-icon"><i class="fas fa-cloud-upload-alt"></i></div>
          <p><strong>Drag &amp; drop</strong> or click to upload</p>
          <p class="sm">PDF, DOCX, or TXT · Maximum 16MB</p>
          <div class="fn" id="sFN"></div>
          <input type="file" id="sF" accept=".pdf,.docx,.txt" onchange="hfs(this,'sFN','sRT')">
        </div>
        <div class="odv">or type / paste below</div>
        <div class="fg">
          <label>Resume text</label>
          <textarea id="sRT" placeholder="Paste your resume content here..."></textarea>
        </div>
      </div>
    </div>
  </div>
  <div class="btn-ctr">
    <button class="btn bp" onclick="doScreen()">
      <i class="fas fa-search"></i> Analyze Resume
    </button>
  </div>

  <!-- Results -->
  <div class="rs" id="sRes">
    <div id="sRC"></div>
  </div>
</div>

<!-- ══════════════ TAB 2: CORRECTIONS ══════════════ -->
<div class="tc" id="tab-correct">
  <div class="ibox inf">
    <i class="fas fa-info-circle"></i>
    <span>Get detailed grammar fixes, formatting improvements, stronger action verbs, quantification suggestions, and content rewrites tailored to your target job.</span>
  </div>
  <div class="g2">
    <div>
      <div class="cd">
        <div class="ct"><div class="icon"><i class="fas fa-briefcase"></i></div> Job Description <span style="font-size:.78rem;color:var(--t3);margin-left:8px">(optional)</span></div>
        <div class="fg">
          <label>For context-aware corrections</label>
          <textarea id="cJD" placeholder="Paste the job description for targeted improvements aligned to the role..."></textarea>
        </div>
      </div>
    </div>
    <div>
      <div class="cd">
        <div class="ct"><div class="icon"><i class="fas fa-file-alt"></i></div> Resume to Review</div>
        <div class="fu">
          <div class="fu-icon"><i class="fas fa-cloud-upload-alt"></i></div>
          <p>Upload your resume file</p>
          <p class="sm">PDF, DOCX, or TXT</p>
          <div class="fn" id="cFN"></div>
          <input type="file" id="cF" accept=".pdf,.docx,.txt" onchange="hfs(this,'cFN','cRT')">
        </div>
        <div class="odv">or paste text</div>
        <div class="fg">
          <label>Resume text <span class="rq">*</span></label>
          <textarea id="cRT" placeholder="Paste your resume here..."></textarea>
        </div>
      </div>
    </div>
  </div>
  <div class="btn-ctr">
    <button class="btn bp" onclick="doCorrect()">
      <i class="fas fa-magic"></i> Get Corrections
    </button>
  </div>
  <div class="rs" id="cRes">
    <div class="cd"><div class="ct"><div class="icon"><i class="fas fa-check-double"></i></div> Corrections &amp; Improvements</div><div id="cRC"></div></div>
  </div>
</div>

<!-- ══════════════ TAB 3: GENERATE ══════════════ -->
<div class="tc" id="tab-generate">
  <div class="ibox inf">
    <i class="fas fa-magic"></i>
    <span>Fill in your details below. AI will generate a complete, professionally formatted, ATS-optimized resume specifically tailored to your target job description.</span>
  </div>

  <div class="cd">
    <div class="ct"><div class="icon"><i class="fas fa-briefcase"></i></div> Target Job Description <span class="rq">*</span></div>
    <div class="fg">
      <textarea class="lg" id="gJD" placeholder="Paste the job description you're applying for. The resume will be crafted to match exactly this role..."></textarea>
    </div>
  </div>

  <div class="cd">
    <div class="ct"><div class="icon"><i class="fas fa-user"></i></div> Personal Information</div>
    <div class="ir">
      <div class="fg"><label>Full name</label><input type="text" id="gN" placeholder="e.g. John Smith"></div>
      <div class="fg"><label>Email address</label><input type="email" id="gE" placeholder="john@example.com"></div>
    </div>
    <div class="ir">
      <div class="fg"><label>Phone number</label><input type="tel" id="gP" placeholder="+1 (555) 123-4567"></div>
      <div class="fg"><label>City, Country</label><input type="text" id="gL" placeholder="New York, USA"></div>
    </div>
    <div class="ir">
      <div class="fg"><label>LinkedIn URL</label><input type="text" id="gLi" placeholder="linkedin.com/in/johnsmith"></div>
      <div class="fg"><label>GitHub / Portfolio</label><input type="text" id="gGit" placeholder="github.com/johnsmith"></div>
    </div>
    <div class="ir">
      <div class="fg"><label>Current job title &amp; company</label><input type="text" id="gCR" placeholder="Senior Engineer at Google"></div>
      <div class="fg"><label>Years of experience</label><input type="text" id="gYE" placeholder="e.g. 7"></div>
    </div>
    <div class="fg"><label>Key technical skills <span style="color:var(--t3);font-weight:400;text-transform:none">(comma separated)</span></label>
      <input type="text" id="gSk" placeholder="Python, React, Node.js, AWS, Docker, PostgreSQL, Machine Learning..."></div>
    <div class="ir">
      <div class="fg"><label>Highest education</label><input type="text" id="gEd" placeholder="B.S. Computer Science, MIT, 2018"></div>
      <div class="fg"><label>Certifications</label><input type="text" id="gCt" placeholder="AWS Solutions Architect, PMP..."></div>
    </div>
    <div class="ir">
      <div class="fg">
        <label>Resume style</label>
        <select id="gSt">
          <option value="professional">Professional (Classic)</option>
          <option value="modern">Modern (Bold & Creative)</option>
          <option value="minimal">Minimal (Clean & Simple)</option>
          <option value="executive">Executive (C-Suite Level)</option>
          <option value="technical">Technical (Engineering Focused)</option>
        </select>
      </div>
      <div class="fg"><label>Industry / Domain</label><input type="text" id="gInd" placeholder="Technology, Finance, Healthcare..."></div>
    </div>
  </div>

  <div class="cd">
    <div class="ct"><div class="icon"><i class="fas fa-file-import"></i></div> Existing Resume <span style="font-size:.78rem;color:var(--t3);margin-left:8px">(optional)</span></div>
    <div class="fg">
      <label>Paste to improve &amp; tailor</label>
      <textarea id="gER" placeholder="If you have a current resume, paste it here. AI will improve and tailor it for the target role..."></textarea>
    </div>
  </div>

  <div class="btn-ctr">
    <button class="btn bg2" onclick="doGenerate()">
      <i class="fas fa-wand-magic-sparkles"></i> Generate My Resume
    </button>
  </div>

  <div class="rs" id="gRes">
    <div class="cd">
      <div class="ct"><div class="icon"><i class="fas fa-file-alt"></i></div> Generated Resume</div>
      <div class="btnr" style="margin-bottom:20px">
        <button class="btn bs" onclick="copyRT()"><i class="fas fa-copy"></i> Copy Plain Text</button>
        <button class="btn bs" onclick="printR()"><i class="fas fa-print"></i> Print / Save PDF</button>
        <button class="btn bs" onclick="downloadTxt()"><i class="fas fa-download"></i> Download TXT</button>
      </div>
      <div id="gRC"></div>
    </div>
  </div>
</div>

<!-- ══════════════ TAB 4: COMPARE ══════════════ -->
<div class="tc" id="tab-compare">
  <div class="ibox inf">
    <i class="fas fa-balance-scale"></i>
    <span>Compare 2–3 candidate resumes side-by-side against the same job description. Get ranked scores, detailed strengths/weaknesses, and an AI-powered hiring recommendation.</span>
  </div>

  <div class="cd">
    <div class="ct"><div class="icon"><i class="fas fa-briefcase"></i></div> Job Description <span class="rq">*</span></div>
    <div class="fg">
      <textarea class="lg" id="xJD" placeholder="Paste the job description to compare all candidates against..."></textarea>
    </div>
  </div>

  <div class="cd">
    <div class="ct"><div class="icon"><i class="fas fa-users"></i></div> Candidate Resumes</div>
    <div class="g2">
      <div class="fg"><label>Candidate 1 <span class="rq">*</span></label><textarea id="xR1" placeholder="Paste Candidate 1 resume text..."></textarea></div>
      <div class="fg"><label>Candidate 2 <span class="rq">*</span></label><textarea id="xR2" placeholder="Paste Candidate 2 resume text..."></textarea></div>
    </div>
    <div id="xR3G" style="display:none">
      <div class="fg"><label>Candidate 3</label><textarea id="xR3" placeholder="Paste Candidate 3 resume text..."></textarea></div>
    </div>
    <button class="btn bs" id="addR3" onclick="togR3()"><i class="fas fa-user-plus"></i> Add 3rd Candidate</button>
  </div>

  <div class="btn-ctr">
    <button class="btn bp" onclick="doCompare()"><i class="fas fa-balance-scale"></i> Compare Candidates</button>
  </div>

  <div class="rs" id="xRes">
    <div class="cd"><div class="ct"><div class="icon"><i class="fas fa-trophy"></i></div> Comparison Results</div><div id="xRC"></div></div>
  </div>
</div>

</div><!-- .wrap -->

<script>
'use strict';
// ── Tab ───────────────────────────────────────────
function stab(n,el){
  document.querySelectorAll('.tc').forEach(t=>t.classList.remove('on'));
  document.querySelectorAll('.tb').forEach(b=>b.classList.remove('on'));
  document.getElementById('tab-'+n).classList.add('on');
  el.classList.add('on');
}

// ── File Upload ───────────────────────────────────
function hfs(inp,fnId,taId){
  const f=inp.files[0]; if(!f) return;
  const el=document.getElementById(fnId);
  el.textContent='✓ '+f.name+' ('+(f.size/1024).toFixed(0)+' KB)';
  el.classList.add('show');
  const fd=new FormData(); fd.append('resume_file',f);
  toast('⏳ Extracting text from '+f.name,'warn');
  fetch('/api/extract-text',{method:'POST',body:fd})
    .then(r=>r.json()).then(d=>{
      if(d.success){
        document.getElementById(taId).value=d.text;
        toast('✓ Extracted '+d.word_count+' words');
      } else toast('Extract failed: '+(d.error||'Unknown'),'err');
    }).catch(e=>toast('Upload error: '+e.message,'err'));
}

// ── Toast ─────────────────────────────────────────
function toast(msg,t){
  t=t||'ok';
  const el=document.getElementById('toast');
  const icons={ok:'<i class="fas fa-check-circle"></i>',err:'<i class="fas fa-times-circle"></i>',warn:'<i class="fas fa-exclamation-triangle"></i>'};
  el.innerHTML=(icons[t]||icons.ok)+' '+msg;
  el.className='toast sh t'+t;
  clearTimeout(el._t);
  el._t=setTimeout(()=>el.classList.remove('sh'),5000);
}

// ── Loading ───────────────────────────────────────
function showL(t,s){
  document.getElementById('loT').textContent=t||'Analyzing...';
  document.getElementById('loS').textContent=s||'Powered by Groq · LLaMA 3.3';
  document.getElementById('lo').classList.add('sh');
}
function hideL(){document.getElementById('lo').classList.remove('sh');}

// ── Score helpers ─────────────────────────────────
function scCls(s){s=parseInt(s)||0;return s>=80?'se':s>=60?'sg':s>=40?'sm2':'spr';}
function scColor(s){s=parseInt(s)||0;return s>=80?'#10b981':s>=60?'#0ea5e9':s>=40?'#f59e0b':'#ef4444';}
function prCls(s){s=parseInt(s)||0;return s>=70?'grn':s>=40?'ylw':'rd';}
function fitCls(l){const m={excellent:'fit-excellent',good:'fit-good',moderate:'fit-moderate',low:'fit-low',poor:'fit-poor'};return m[(l||'').toLowerCase()]||'fit-moderate';}
function probCls(p){const m={high:'prob-high',medium:'prob-medium',low:'prob-low'};return m[(p||'').toLowerCase()]||'prob-medium';}
function cap(s){return(s||'').replace(/_/g,' ').replace(/\b\w/g,c=>c.toUpperCase());}

function mkPB(label,val,cls){
  val=parseInt(val)||0; cls=cls||prCls(val);
  return `<div class="pi">
    <div class="ph"><span class="ph-label">${label}</span><span class="ph-val">${val}%</span></div>
    <div class="pbar"><div class="pfill ${cls}" style="width:${val}%"></div></div>
  </div>`;
}

function mkRing(score,size){
  size=size||160; score=parseInt(score)||0;
  const r=size/2-10; const circ=2*Math.PI*r;
  const dash=circ*(score/100); const gap=circ-dash;
  const col=scColor(score);
  return `<svg width="${size}" height="${size}" viewBox="0 0 ${size} ${size}">
    <circle cx="${size/2}" cy="${size/2}" r="${r}" fill="none" stroke="rgba(30,41,59,.6)" stroke-width="10"/>
    <circle cx="${size/2}" cy="${size/2}" r="${r}" fill="none" stroke="${col}" stroke-width="10"
      stroke-dasharray="${dash} ${gap}" stroke-linecap="round"
      style="filter:drop-shadow(0 0 6px ${col}40)"/>
  </svg>`;
}

// ── Error render ──────────────────────────────────
function renderErr(id,msg){
  document.getElementById(id).innerHTML=`
    <div class="err-box">
      <strong><i class="fas fa-exclamation-triangle"></i> Error</strong><br><br>
      ${msg}<br><br>
      <small>Check your Groq API key and visit <a href="https://console.groq.com/docs/models" target="_blank">console.groq.com/docs/models</a> for available models.</small>
    </div>`;
}

// ══════════════════════════════════════════════════
// TAB 1 — SCREEN
// ══════════════════════════════════════════════════
async function doScreen(){
  const jd=document.getElementById('sJD').value.trim();
  const rt=document.getElementById('sRT').value.trim();
  const fi=document.getElementById('sF');
  if(!jd){toast('Please enter a job description','err');return;}
  if(!rt&&!fi.files.length){toast('Please provide your resume','err');return;}
  showL('Screening your resume...','Analyzing keywords, ATS compatibility & fit score');
  const fd=new FormData();
  fd.append('job_description',jd); fd.append('resume_text',rt);
  if(fi.files.length) fd.append('resume_file',fi.files[0]);
  try{
    const r=await fetch('/api/screen-resume',{method:'POST',body:fd});
    const d=await r.json();
    if(d.success){renderScreen(d.data);toast('✓ Analysis complete!');}
    else{renderErr('sRC',d.error||'Screening failed');document.getElementById('sRes').classList.add('sh');toast(d.error||'Failed','err');}
  }catch(e){renderErr('sRC','Network error: '+e.message);document.getElementById('sRes').classList.add('sh');toast('Error: '+e.message,'err');}
  finally{hideL();}
}

function renderScreen(d){
  if(!d||d.raw_response){renderErr('sRC','Could not parse AI response. Please try again.');document.getElementById('sRes').classList.add('sh');return;}
  const os=parseInt(d.overall_score)||0;
  const col=scColor(os);
  let h='';

  // ── Score Hero ──
  h+=`<div class="score-hero">
    <div class="score-ring-wrap">
      <div class="score-ring">
        ${mkRing(os,160)}
        <div class="score-ring-inner">
          <div class="score-val" style="color:${col}">${os}</div>
          <div class="score-lbl">Overall Score</div>
        </div>
      </div>
    </div>
    <div><span class="fit-chip ${fitCls(d.fit_level)}">
      ${d.fit_level==='EXCELLENT'?'⭐':d.fit_level==='GOOD'?'✅':d.fit_level==='MODERATE'?'⚡':d.fit_level==='LOW'?'⚠️':'❌'}
      ${d.fit_level||'N/A'} FIT
    </span></div>
    <p class="hero-summary">${d.summary||''}</p>
    <div><span class="prob-badge ${probCls(d.interview_probability)}">
      <i class="fas fa-handshake"></i> Interview Probability: ${d.interview_probability||'N/A'}
    </span></div>
  </div>`;

  // ── Stat Cards ──
  const stats=[
    {l:'ATS Score',v:d.ats_compatibility_score,c:'c-purple'},
    {l:'Keyword Match',v:d.keyword_match_score,c:'c-blue'},
    {l:'Experience',v:d.experience_relevance_score,c:'c-green'},
    {l:'Skills Match',v:d.skills_match_score,c:'c-orange'},
  ];
  h+=`<div class="stat-grid">`;
  stats.forEach(st=>{
    const v=parseInt(st.v)||0; const cl=prCls(v);
    const col2=cl==='grn'?'#10b981':cl==='ylw'?'#f59e0b':'#ef4444';
    h+=`<div class="stat-card ${st.c}">
      <div class="stat-v" style="color:${col2}">${v}</div>
      <div class="stat-l">${st.l}</div>
      <div class="stat-bar"><div class="stat-bar-fill" style="width:${v}%;background:${col2}"></div></div>
    </div>`;
  });
  h+=`</div>`;

  // ── Score Breakdown + Sections ──
  h+=`<div class="g2" style="margin-bottom:20px">`;
  h+=`<div class="cd"><div class="sec-hdr"><div class="ico" style="background:rgba(124,58,237,.15);color:var(--pl)"><i class="fas fa-chart-line"></i></div><h3>Score Breakdown</h3></div>`;
  h+=mkPB('ATS Compatibility',d.ats_compatibility_score,'pur');
  h+=mkPB('Keyword Match',d.keyword_match_score,'blu');
  h+=mkPB('Experience Relevance',d.experience_relevance_score,'grn');
  h+=mkPB('Skills Match',d.skills_match_score,'grn');
  h+=mkPB('Education Match',d.education_match_score,'blu');
  h+=mkPB('Formatting',d.formatting_score,'pur');
  h+=`</div>`;

  h+=`<div class="cd"><div class="sec-hdr"><div class="ico" style="background:rgba(14,165,233,.15);color:#38bdf8"><i class="fas fa-th-large"></i></div><h3>Section Scores</h3></div>`;
  if(d.section_scores&&typeof d.section_scores==='object'){
    Object.entries(d.section_scores).forEach(([k,v])=>h+=mkPB(cap(k),v));
  }else h+=`<p style="color:var(--t3);font-size:.9rem">No section data</p>`;
  h+=`</div></div>`;

  // ── Keywords ──
  const mk=d.matched_keywords||[]; const ms=d.missing_keywords||[];
  h+=`<div class="g2" style="margin-bottom:20px">
    <div class="cd">
      <div class="sec-hdr"><div class="ico" style="background:rgba(16,185,129,.15);color:#34d399"><i class="fas fa-check-circle"></i></div><h3>Matched Keywords (${mk.length})</h3></div>
      <div class="tags">${mk.map(k=>`<span class="tag tg">${k}</span>`).join('')||'<span style="color:var(--t3);font-size:.88rem">None found</span>'}
      </div>
    </div>
    <div class="cd">
      <div class="sec-hdr"><div class="ico" style="background:rgba(239,68,68,.15);color:#f87171"><i class="fas fa-times-circle"></i></div><h3>Missing Keywords (${ms.length})</h3></div>
      <div class="tags">${ms.map(k=>`<span class="tag tr">${k}</span>`).join('')||'<span style="color:var(--t3);font-size:.88rem">None — great!</span>'}
      </div>
    </div>
  </div>`;

  // ── Strengths & Weaknesses ──
  h+=`<div class="g2" style="margin-bottom:20px">
    <div class="cd">
      <div class="sec-hdr"><div class="ico" style="background:rgba(16,185,129,.15);color:#34d399"><i class="fas fa-thumbs-up"></i></div><h3>Strengths</h3></div>
      <ul class="rl grl">${(d.strengths||[]).map(s=>`<li>✓ ${s}</li>`).join('')||'<li>No strengths identified</li>'}</ul>
    </div>
    <div class="cd">
      <div class="sec-hdr"><div class="ico" style="background:rgba(239,68,68,.15);color:#f87171"><i class="fas fa-exclamation-circle"></i></div><h3>Areas to Improve</h3></div>
      <ul class="rl rrl">${(d.weaknesses||[]).map(w=>`<li>⚠ ${w}</li>`).join('')||'<li>No major weaknesses</li>'}</ul>
    </div>
  </div>`;

  // ── Skills Gap ──
  if((d.skills_gap||[]).length){
    h+=`<div class="cd" style="margin-bottom:20px">
      <div class="sec-hdr"><div class="ico" style="background:rgba(245,158,11,.15);color:#fbbf24"><i class="fas fa-puzzle-piece"></i></div><h3>Skills Gap to Address</h3></div>
      <div class="tags">${d.skills_gap.map(s=>`<span class="tag ty">+ ${s}</span>`).join('')}</div>
    </div>`;
  }

  // ── Experience Analysis ──
  if(d.experience_analysis){
    h+=`<div class="cd" style="margin-bottom:20px">
      <div class="sec-hdr"><div class="ico" style="background:rgba(124,58,237,.15);color:var(--pl)"><i class="fas fa-briefcase"></i></div><h3>Experience Analysis</h3></div>
      <p style="color:var(--t2);line-height:1.8;font-size:.92rem">${d.experience_analysis}</p>
    </div>`;
  }

  // ── ATS Issues ──
  if((d.ats_issues||[]).length){
    h+=`<div class="cd" style="margin-bottom:20px">
      <div class="sec-hdr"><div class="ico" style="background:rgba(245,158,11,.15);color:#fbbf24"><i class="fas fa-robot"></i></div><h3>ATS Compatibility Issues</h3></div>
      <ul class="rl yrl">${d.ats_issues.map(i=>`<li><i class="fas fa-exclamation-triangle"></i> ${i}</li>`).join('')}</ul>
    </div>`;
  }

  // ── Recommendations ──
  h+=`<div class="cd">
    <div class="sec-hdr"><div class="ico" style="background:rgba(245,158,11,.15);color:#fbbf24"><i class="fas fa-lightbulb"></i></div><h3>Top Recommendations</h3></div>
    <ul class="rl">${(d.recommendations||[]).map((r,i)=>`<li><strong style="color:var(--pl)">${i+1}.</strong> ${r}</li>`).join('')||'<li>No recommendations</li>'}</ul>
  </div>`;

  document.getElementById('sRC').innerHTML=h;
  document.getElementById('sRes').classList.add('sh');
  document.getElementById('sRes').scrollIntoView({behavior:'smooth',block:'start'});

  // Animate bars after render
  setTimeout(()=>{document.querySelectorAll('.pfill,.stat-bar-fill').forEach(b=>{const w=b.style.width;b.style.width='0';requestAnimationFrame(()=>b.style.width=w);});},100);
}

// ══════════════════════════════════════════════════
// TAB 2 — CORRECTIONS
// ══════════════════════════════════════════════════
async function doCorrect(){
  const jd=document.getElementById('cJD').value.trim();
  const rt=document.getElementById('cRT').value.trim();
  const fi=document.getElementById('cF');
  if(!rt&&!fi.files.length){toast('Please provide your resume','err');return;}
  showL('Analyzing for corrections...','Checking grammar, formatting, action verbs & content');
  const fd=new FormData();
  fd.append('job_description',jd); fd.append('resume_text',rt);
  if(fi.files.length) fd.append('resume_file',fi.files[0]);
  try{
    const r=await fetch('/api/get-corrections',{method:'POST',body:fd});
    const d=await r.json();
    if(d.success){renderCorrections(d.data);toast('✓ Corrections ready!');}
    else{renderErr('cRC',d.error||'Failed');document.getElementById('cRes').classList.add('sh');toast(d.error||'Failed','err');}
  }catch(e){renderErr('cRC','Network error: '+e.message);document.getElementById('cRes').classList.add('sh');toast('Error: '+e.message,'err');}
  finally{hideL();}
}

function renderCorrections(d){
  if(!d||d.raw_response){renderErr('cRC','Could not parse response. Please try again.');document.getElementById('cRes').classList.add('sh');return;}
  let h='';

  if(d.overall_feedback){
    h+=`<div class="ibox inf" style="margin-bottom:20px">
      <i class="fas fa-comment-dots"></i>
      <div><strong style="display:block;margin-bottom:5px">Overall Feedback</strong>${d.overall_feedback}</div>
    </div>`;
  }

  if((d.priority_fixes||[]).length){
    h+=`<div class="cd" style="margin-bottom:20px">
      <div class="sec-hdr"><div class="ico" style="background:rgba(239,68,68,.15);color:#f87171"><i class="fas fa-fire"></i></div><h3>Priority Fixes</h3></div>
      <ul class="rl rrl">${d.priority_fixes.map(f=>`<li>⚡ <strong>${f}</strong></li>`).join('')}</ul>
    </div>`;
  }

  h+=`<div class="g2" style="margin-bottom:20px">`;
  if(d.word_count_analysis){const w=d.word_count_analysis;
    h+=`<div class="cd" style="margin-bottom:0">
      <div class="sec-hdr"><div class="ico" style="background:rgba(14,165,233,.15);color:#38bdf8"><i class="fas fa-align-left"></i></div><h3>Word Count</h3></div>
      <p style="color:var(--t2);margin-bottom:8px">Current estimate: <strong style="color:var(--t1)">${w.current_estimated||'N/A'}</strong></p>
      <p style="color:var(--t2);margin-bottom:12px">Recommended: <strong style="color:var(--t1)">${w.recommended||'N/A'}</strong></p>
      <span class="tag ${w.verdict==='appropriate'?'tg':'ty'}">${(w.verdict||'N/A').toUpperCase()}</span>
    </div>`;
  }
  if(d.tone_analysis){const ta=d.tone_analysis;
    h+=`<div class="cd" style="margin-bottom:0">
      <div class="sec-hdr"><div class="ico" style="background:rgba(124,58,237,.15);color:var(--pl)"><i class="fas fa-comment"></i></div><h3>Tone Analysis</h3></div>
      <p style="color:var(--t2);margin-bottom:10px">Current: <span class="tag tp">${ta.current_tone||'N/A'}</span></p>
      <p style="color:var(--t2);font-size:.9rem;line-height:1.65">${ta.recommended_changes||''}</p>
    </div>`;
  }
  h+=`</div>`;

  if((d.grammar_issues||[]).length){
    h+=`<div class="cd" style="margin-bottom:20px">
      <div class="sec-hdr"><div class="ico" style="background:rgba(239,68,68,.15);color:#f87171"><i class="fas fa-spell-check"></i></div><h3>Grammar Issues (${d.grammar_issues.length})</h3></div>
      ${d.grammar_issues.map(g=>`<div class="cc">
        ${g.original?`<div class="og"><i class="fas fa-times"></i> ${g.original}</div>`:''}
        <div class="cr"><i class="fas fa-check"></i> ${g.corrected||''}</div>
        ${g.explanation?`<div class="ex">${g.explanation}</div>`:''}
      </div>`).join('')}
    </div>`;
  }

  if((d.formatting_issues||[]).length){
    h+=`<div class="cd" style="margin-bottom:20px">
      <div class="sec-hdr"><div class="ico" style="background:rgba(245,158,11,.15);color:#fbbf24"><i class="fas fa-align-left"></i></div><h3>Formatting Issues</h3></div>
      ${d.formatting_issues.map(f=>{
        const pr=(f.priority||'MEDIUM').toUpperCase();
        const pc=pr==='HIGH'?'prh':pr==='LOW'?'prl':'prm';
        return `<div class="cc">
          <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:10px">
            <strong style="color:var(--t1)">${f.issue||''}</strong>
            <span class="prb ${pc}">${pr}</span>
          </div>
          <p style="color:var(--t2);font-size:.9rem">→ ${f.fix||''}</p>
        </div>`;
      }).join('')}
    </div>`;
  }

  if((d.content_improvements||[]).length){
    h+=`<div class="cd" style="margin-bottom:20px">
      <div class="sec-hdr"><div class="ico" style="background:rgba(124,58,237,.15);color:var(--pl)"><i class="fas fa-pencil-alt"></i></div><h3>Content Improvements</h3></div>
      ${d.content_improvements.map(ci=>`<div class="cc">
        <div style="margin-bottom:10px"><span class="tag tp">${ci.section||''}</span></div>
        ${ci.current?`<div class="og">${ci.current}</div>`:''}
        <div class="cr">${ci.suggested||''}</div>
        ${ci.reason?`<div class="ex">💡 ${ci.reason}</div>`:''}
      </div>`).join('')}
    </div>`;
  }

  if((d.action_verb_suggestions||[]).length){
    h+=`<div class="cd" style="margin-bottom:20px">
      <div class="sec-hdr"><div class="ico" style="background:rgba(16,185,129,.15);color:#34d399"><i class="fas fa-bolt"></i></div><h3>Stronger Action Verbs</h3></div>
      ${d.action_verb_suggestions.map(a=>`<div class="cc" style="display:flex;align-items:center;flex-wrap:wrap;gap:8px">
        <span class="tag tr">${a.weak_verb||''}</span>
        <span style="color:var(--t3);font-size:1.3rem">→</span>
        ${(a.strong_alternatives||[]).map(v=>`<span class="tag tg">${v}</span>`).join('')}
      </div>`).join('')}
    </div>`;
  }

  if((d.quantification_opportunities||[]).length){
    h+=`<div class="cd">
      <div class="sec-hdr"><div class="ico" style="background:rgba(14,165,233,.15);color:#38bdf8"><i class="fas fa-percentage"></i></div><h3>Add Metrics &amp; Numbers</h3></div>
      ${d.quantification_opportunities.map(q=>`<div class="cc">
        ${q.bullet_point?`<div class="og">${q.bullet_point}</div>`:''}
        <div class="cr">${q.improved||''}</div>
        ${q.tip?`<div class="ex">💡 ${q.tip}</div>`:''}
      </div>`).join('')}
    </div>`;
  }

  document.getElementById('cRC').innerHTML=h;
  document.getElementById('cRes').classList.add('sh');
  document.getElementById('cRes').scrollIntoView({behavior:'smooth',block:'start'});
}

// ══════════════════════════════════════════════════
// TAB 3 — GENERATE
// ══════════════════════════════════════════════════
let genRT='';
async function doGenerate(){
  const jd=document.getElementById('gJD').value.trim();
  if(!jd){toast('Please enter a job description','err');return;}
  showL('Generating your tailored resume...','Crafting ATS-optimized content for this specific role');
  const payload={
    job_description:jd,
    existing_resume:document.getElementById('gER').value.trim(),
    style:document.getElementById('gSt').value,
    user_info:{
      name:document.getElementById('gN').value.trim(),
      email:document.getElementById('gE').value.trim(),
      phone:document.getElementById('gP').value.trim(),
      location:document.getElementById('gL').value.trim(),
      linkedin:document.getElementById('gLi').value.trim(),
      github:document.getElementById('gGit').value.trim(),
      current_role:document.getElementById('gCR').value.trim(),
      experience_years:document.getElementById('gYE').value.trim(),
      skills:document.getElementById('gSk').value.trim(),
      education:document.getElementById('gEd').value.trim(),
      certifications:document.getElementById('gCt').value.trim(),
      industry:document.getElementById('gInd').value.trim()
    }
  };
  try{
    const r=await fetch('/api/generate-resume',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(payload)});
    const d=await r.json();
    if(d.success){renderGen(d.data);toast('✓ Resume generated!');}
    else{renderErr('gRC',d.error||'Generation failed');document.getElementById('gRes').classList.add('sh');toast(d.error||'Failed','err');}
  }catch(e){renderErr('gRC','Network error: '+e.message);document.getElementById('gRes').classList.add('sh');toast('Error: '+e.message,'err');}
  finally{hideL();}
}

function renderGen(d){
  if(!d){renderErr('gRC','No data returned');return;}
  genRT=d.resume_text||'';
  let h='';
  if(d.tailoring_notes){
    h+=`<div class="ibox ok" style="margin-bottom:18px">
      <i class="fas fa-check-circle"></i>
      <div><strong>Tailoring Notes:</strong><br>${d.tailoring_notes}</div>
    </div>`;
  }
  if((d.keyword_optimization||[]).length){
    h+=`<div style="margin-bottom:18px"><p style="color:var(--t2);font-size:.85rem;font-weight:600;text-transform:uppercase;letter-spacing:.5px;margin-bottom:10px">Keywords Integrated</p>
      <div class="tags">${d.keyword_optimization.map(k=>`<span class="tag tg">${k}</span>`).join('')}</div>
    </div>`;
  }
  if(d.resume_html){
    h+=`<p style="color:var(--t2);font-size:.85rem;font-weight:600;text-transform:uppercase;letter-spacing:.5px;margin-bottom:10px"><i class="fas fa-eye"></i> Preview</p>
    <div class="rpv" id="rpv">${d.resume_html}</div>`;
  }
  if(d.resume_text){
    h+=`<div style="margin-top:24px">
      <p style="color:var(--t2);font-size:.85rem;font-weight:600;text-transform:uppercase;letter-spacing:.5px;margin-bottom:10px"><i class="fas fa-file-alt"></i> Plain Text Version</p>
      <textarea style="width:100%;min-height:480px;font-family:monospace;font-size:.82rem;line-height:1.65" readonly>${d.resume_text}</textarea>
    </div>`;
  }
  document.getElementById('gRC').innerHTML=h;
  document.getElementById('gRes').classList.add('sh');
  document.getElementById('gRes').scrollIntoView({behavior:'smooth',block:'start'});
}

function copyRT(){
  if(!genRT){toast('Generate a resume first','err');return;}
  navigator.clipboard.writeText(genRT).then(()=>toast('✓ Copied to clipboard!')).catch(()=>toast('Copy failed','err'));
}

function downloadTxt(){
  if(!genRT){toast('Generate a resume first','err');return;}
  const a=document.createElement('a');
  a.href='data:text/plain;charset=utf-8,'+encodeURIComponent(genRT);
  a.download='resume.txt'; a.click();
  toast('✓ Downloading resume.txt');
}

function printR(){
  const p=document.getElementById('rpv');
  if(!p){toast('Generate a resume first','err');return;}
  const w=window.open('','_blank');
  w.document.write(`<!DOCTYPE html><html><head><title>Resume</title><style>
    body{font-family:Arial,sans-serif;margin:40px auto;max-width:820px;color:#1a1a1a;line-height:1.7}
    h1{font-size:1.8rem;color:#1a1a2e;border-bottom:3px solid #7c3aed;padding-bottom:12px}
    h2{font-size:1.1rem;color:#1a1a2e;border-bottom:1px solid #e2e8f0;padding-bottom:5px;margin-top:20px}
    ul{margin-left:20px}li{margin-bottom:5px}a{color:#7c3aed}
    @media print{@page{margin:18mm}body{margin:0}}
  </style></head><body>${p.innerHTML}</body></html>`);
  w.document.close(); setTimeout(()=>w.print(),500);
}

// ══════════════════════════════════════════════════
// TAB 4 — COMPARE
// ══════════════════════════════════════════════════
let sh3=false;
function togR3(){
  sh3=!sh3;
  document.getElementById('xR3G').style.display=sh3?'block':'none';
  document.getElementById('addR3').innerHTML=sh3?'<i class="fas fa-user-minus"></i> Remove 3rd Candidate':'<i class="fas fa-user-plus"></i> Add 3rd Candidate';
}

async function doCompare(){
  const jd=document.getElementById('xJD').value.trim();
  const r1=document.getElementById('xR1').value.trim();
  const r2=document.getElementById('xR2').value.trim();
  const r3=document.getElementById('xR3').value.trim();
  if(!jd){toast('Please enter a job description','err');return;}
  if(!r1||!r2){toast('Provide at least 2 candidate resumes','err');return;}
  showL('Comparing candidates...','Ranking and scoring each resume against the job');
  const fd=new FormData();
  fd.append('job_description',jd);
  fd.append('resume_text_1',r1); fd.append('resume_text_2',r2);
  if(r3) fd.append('resume_text_3',r3);
  try{
    const r=await fetch('/api/compare-resumes',{method:'POST',body:fd});
    const d=await r.json();
    if(d.success){renderCompare(d.data);toast('✓ Comparison complete!');}
    else{renderErr('xRC',d.error||'Failed');document.getElementById('xRes').classList.add('sh');toast(d.error||'Failed','err');}
  }catch(e){renderErr('xRC','Network error: '+e.message);document.getElementById('xRes').classList.add('sh');toast('Error: '+e.message,'err');}
  finally{hideL();}
}

function renderCompare(d){
  if(!d||d.raw_response){renderErr('xRC','Could not parse results. Please try again.');document.getElementById('xRes').classList.add('sh');return;}
  let h='';
  const medals=['🥇','🥈','🥉'];
  const bcolors=['rgba(16,185,129,.4)','rgba(14,165,233,.4)','rgba(245,158,11,.4)'];
  if((d.rankings||[]).length){
    const sorted=[...d.rankings].sort((a,b)=>a.rank-b.rank);
    sorted.forEach((r,i)=>{
      const v=parseInt(r.overall_score)||0; const col=scColor(v);
      h+=`<div class="cmp-card" style="border-left:4px solid ${bcolors[i]||'rgba(124,58,237,.4)'}">
        <div class="cmp-rank">
          <div>
            <div class="cmp-title">${medals[i]||'#'+r.rank} Rank #${r.rank} — ${r.resume_name||'Candidate '+r.rank}</div>
            ${r.fit_summary?`<p style="color:var(--t2);font-size:.88rem;margin-top:5px;line-height:1.55">${r.fit_summary}</p>`:''}
          </div>
          <div class="mini-score ${scCls(v)}" style="color:${col};border:2px solid ${col}">${v}</div>
        </div>
        <div class="g2">
          <div>
            <p style="color:#34d399;font-size:.78rem;font-weight:700;text-transform:uppercase;letter-spacing:.5px;margin-bottom:8px">✓ Strengths</p>
            <ul class="rl grl">${(r.strengths||[]).map(s=>`<li style="font-size:.85rem;padding:9px 13px">${s}</li>`).join('')}</ul>
          </div>
          <div>
            <p style="color:#f87171;font-size:.78rem;font-weight:700;text-transform:uppercase;letter-spacing:.5px;margin-bottom:8px">✗ Weaknesses</p>
            <ul class="rl rrl">${(r.weaknesses||[]).map(w=>`<li style="font-size:.85rem;padding:9px 13px">${w}</li>`).join('')}</ul>
          </div>
        </div>
      </div>`;
    });
  }
  if(d.comparison_summary){
    h+=`<div class="ibox inf" style="margin-top:16px">
      <i class="fas fa-clipboard-list"></i>
      <div><strong style="display:block;margin-bottom:5px">Comparison Summary</strong>${d.comparison_summary}</div>
    </div>`;
  }
  if(d.recommendation){
    h+=`<div class="ibox ok" style="margin-top:10px">
      <i class="fas fa-star"></i>
      <div><strong style="display:block;margin-bottom:5px">Hiring Recommendation</strong>${d.recommendation}</div>
    </div>`;
  }
  document.getElementById('xRC').innerHTML=h;
  document.getElementById('xRes').classList.add('sh');
  document.getElementById('xRes').scrollIntoView({behavior:'smooth',block:'start'});
}

// ── Drag & Drop ───────────────────────────────────
document.querySelectorAll('.fu').forEach(a=>{
  a.addEventListener('dragover',e=>{e.preventDefault();a.classList.add('dg');});
  a.addEventListener('dragleave',()=>a.classList.remove('dg'));
  a.addEventListener('drop',e=>{
    e.preventDefault();a.classList.remove('dg');
    const inp=a.querySelector('input[type=file]');
    if(e.dataTransfer.files.length&&inp){inp.files=e.dataTransfer.files;inp.dispatchEvent(new Event('change'));}
  });
});

// ── Health check ──────────────────────────────────
window.addEventListener('load',()=>{
  fetch('/api/health').then(r=>r.json()).then(d=>{
    const el=document.getElementById('modelBdg');
    if(d.api_configured){
      el.innerHTML=`<i class="fas fa-microchip"></i> ${d.active_model}`;
      el.className='bdg bdg-green';
    } else {
      el.innerHTML='<i class="fas fa-times-circle"></i> API Key Not Set';
      el.className='bdg';
      el.style.cssText='background:rgba(239,68,68,.15);border-color:rgba(239,68,68,.3);color:#f87171';
      toast('⚠️ Set GROQ_API_KEY in resume.py (line ~15)','err');
    }
  }).catch(()=>toast('Cannot reach backend server','err'));
});
</script>
</body>
</html>"""

# ── API Routes ────────────────────────────────────────────────
@app.route("/")
def index(): return Response(HTML, mimetype="text/html")

@app.route("/api/health")
def health():
    active = MODEL_NAME
    if client:
        for m in FALLBACK_MODELS:
            try:
                client.chat.completions.create(model=m,messages=[{"role":"user","content":"OK"}],max_tokens=3)
                active=m; break
            except: continue
    return jsonify({"status":"ok","api_configured":client is not None,"active_model":active,"models":FALLBACK_MODELS,"ts":datetime.now().isoformat()})

@app.route("/api/extract-text", methods=["POST"])
def extract_text():
    try:
        if "resume_file" not in request.files: return jsonify({"error":"No file"}),400
        f=request.files["resume_file"]
        if not f or not f.filename: return jsonify({"error":"No file selected"}),400
        if not allowed_file(f.filename): return jsonify({"error":"Allowed: PDF, DOCX, TXT"}),400
        t=read_file(f)
        if not t.strip(): return jsonify({"error":"Could not extract text"}),400
        return jsonify({"success":True,"text":t,"filename":f.filename,"word_count":len(t.split())})
    except Exception as e:
        traceback.print_exc(); return jsonify({"error":str(e)}),500

@app.route("/api/screen-resume", methods=["POST"])
def screen_resume():
    try:
        jd=request.form.get("job_description","").strip()
        rt=request.form.get("resume_text","").strip()
        if "resume_file" in request.files:
            f=request.files["resume_file"]
            if f and f.filename and allowed_file(f.filename): rt=read_file(f)
        if not jd: return jsonify({"error":"Job description required"}),400
        if not rt: return jsonify({"error":"Resume required"}),400

        prompt=f"""Analyze this resume vs the job description. Return ONLY raw valid JSON (no markdown):
{{
  "overall_score":75,"ats_compatibility_score":80,"keyword_match_score":70,
  "experience_relevance_score":75,"skills_match_score":80,
  "education_match_score":90,"formatting_score":85,
  "summary":"2-3 sentence overall assessment",
  "strengths":["strength 1","strength 2","strength 3"],
  "weaknesses":["weakness 1","weakness 2","weakness 3"],
  "missing_keywords":["kw1","kw2","kw3"],
  "matched_keywords":["kw1","kw2","kw3"],
  "experience_analysis":"Detailed paragraph about experience relevance",
  "skills_gap":["skill1","skill2"],
  "recommendations":["rec1","rec2","rec3","rec4","rec5"],
  "ats_issues":["issue1","issue2"],
  "section_scores":{{"contact_info":90,"professional_summary":75,"work_experience":80,"skills":70,"education":85,"certifications":60,"projects":70}},
  "fit_level":"GOOD",
  "interview_probability":"MEDIUM"
}}
fit_level: EXCELLENT|GOOD|MODERATE|LOW|POOR
interview_probability: HIGH|MEDIUM|LOW

JOB:\n{jd}\n\nRESUME:\n{rt}"""

        res=call_ai(prompt)
        data=parse_json(res)
        defs={"overall_score":0,"ats_compatibility_score":0,"keyword_match_score":0,"experience_relevance_score":0,
              "skills_match_score":0,"education_match_score":0,"formatting_score":0,"summary":"","strengths":[],
              "weaknesses":[],"missing_keywords":[],"matched_keywords":[],"experience_analysis":"","skills_gap":[],
              "recommendations":[],"ats_issues":[],"section_scores":{},"fit_level":"MODERATE","interview_probability":"MEDIUM"}
        for k,v in defs.items():
            if k not in data: data[k]=v
        return jsonify({"success":True,"data":data})
    except Exception as e:
        traceback.print_exc(); return jsonify({"error":str(e)}),500

@app.route("/api/get-corrections", methods=["POST"])
def get_corrections():
    try:
        jd=request.form.get("job_description","").strip()
        rt=request.form.get("resume_text","").strip()
        if "resume_file" in request.files:
            f=request.files["resume_file"]
            if f and f.filename and allowed_file(f.filename): rt=read_file(f)
        if not rt: return jsonify({"error":"Resume required"}),400
        jd_part=f"\nJOB:\n{jd}\n" if jd else ""

        prompt=f"""Review this resume, provide corrections. Return ONLY raw valid JSON:
{{
  "grammar_issues":[{{"original":"text with error","corrected":"fixed text","explanation":"why"}}],
  "formatting_issues":[{{"issue":"description","fix":"how to fix","priority":"HIGH"}}],
  "content_improvements":[{{"section":"Work Experience","current":"weak content","suggested":"strong improved","reason":"why better"}}],
  "action_verb_suggestions":[{{"weak_verb":"managed","strong_alternatives":["spearheaded","orchestrated","directed"]}}],
  "quantification_opportunities":[{{"bullet_point":"current bullet","improved":"with numbers","tip":"add metric"}}],
  "overall_feedback":"Comprehensive feedback paragraph",
  "priority_fixes":["critical fix 1","critical fix 2","critical fix 3"],
  "word_count_analysis":{{"current_estimated":"450","recommended":"400-600","verdict":"appropriate"}},
  "tone_analysis":{{"current_tone":"professional","recommended_changes":"suggestions"}}
}}
{jd_part}
RESUME:\n{rt}"""

        res=call_ai(prompt)
        data=parse_json(res)
        return jsonify({"success":True,"data":data})
    except Exception as e:
        traceback.print_exc(); return jsonify({"error":str(e)}),500

@app.route("/api/generate-resume", methods=["POST"])
def generate_resume():
    try:
        body=request.get_json()
        if not body: return jsonify({"error":"No JSON"}),400
        jd=body.get("job_description","").strip()
        ui=body.get("user_info",{})
        er=body.get("existing_resume","").strip()
        style=body.get("style","professional")
        if not jd: return jsonify({"error":"Job description required"}),400

        info="USER INFO:\n"
        for lbl,key in [("Name","name"),("Email","email"),("Phone","phone"),("Location","location"),
                        ("LinkedIn","linkedin"),("GitHub","github"),("Experience","experience_years"),
                        ("Role","current_role"),("Skills","skills"),("Education","education"),
                        ("Certs","certifications"),("Industry","industry")]:
            v=ui.get(key,"")
            if v: info+=f"- {lbl}: {v}\n"

        er_part=f"\nEXISTING RESUME:\n{er}\n" if er else ""

        prompt=f"""Create a professional ATS-optimized resume. Style: {style}
Return ONLY raw valid JSON:
{{
  "resume_html":"<complete professional HTML resume with inline CSS>",
  "resume_text":"complete plain text resume",
  "resume_sections":{{"contact":{{"name":"","email":"","phone":"","location":"","linkedin":""}},"professional_summary":"","work_experience":[{{"title":"","company":"","duration":"","bullets":["achievement with metric"]}}],"skills":{{"technical":[],"soft":[],"tools":[]}},"education":[{{"degree":"","institution":"","year":""}}],"certifications":[],"projects":[{{"name":"","description":"","technologies":[]}}]}},
  "keyword_optimization":["kw1","kw2","kw3"],
  "tailoring_notes":"How resume was tailored"
}}

resume_html must be a complete styled HTML resume with:
- Professional header with name prominent
- Clean section headers with dividers
- Bullet points with strong action verbs and metrics
- Subtle color accent (use #7c3aed for headings/accents)
- Proper spacing and hierarchy
- All inline CSS

{info}{er_part}
JOB:\n{jd}"""

        res=call_ai(prompt)
        data=parse_json(res)
        return jsonify({"success":True,"data":data})
    except Exception as e:
        traceback.print_exc(); return jsonify({"error":str(e)}),500

@app.route("/api/compare-resumes", methods=["POST"])
def compare_resumes():
    try:
        jd=request.form.get("job_description","").strip()
        resumes=[]
        for i in range(1,6):
            t=request.form.get(f"resume_text_{i}","").strip()
            if t: resumes.append({"name":f"Candidate {i}","text":t})
        if not jd: return jsonify({"error":"Job description required"}),400
        if len(resumes)<2: return jsonify({"error":"At least 2 resumes needed"}),400

        block="".join(f"\n--- RESUME {i+1} ({r['name']}) ---\n{r['text']}\n" for i,r in enumerate(resumes))

        prompt=f"""Compare {len(resumes)} resumes vs the job. Return ONLY raw valid JSON:
{{
  "rankings":[{{"rank":1,"resume_name":"Candidate 1","overall_score":85,"strengths":["s1","s2"],"weaknesses":["w1"],"fit_summary":"1-2 sentence summary"}}],
  "comparison_summary":"Comprehensive comparison paragraph",
  "recommendation":"Clear hiring recommendation with reasoning",
  "detailed_comparison":{{"skills_match":[{{"resume":"Candidate 1","score":85,"matched":[],"missing":[]}}],"experience_relevance":[{{"resume":"Candidate 1","score":80,"analysis":"brief"}}]}}
}}

JOB:\n{jd}\n{block}"""

        res=call_ai(prompt)
        data=parse_json(res)
        return jsonify({"success":True,"data":data})
    except Exception as e:
        traceback.print_exc(); return jsonify({"error":str(e)}),500

# ── Run ───────────────────────────────────────────────────────
if __name__=="__main__":
    print("\n"+"="*60)
    print("  🤖  AI Resume Screener Pro")
    print("  📍  http://localhost:5000")
    print(f"  🔑  Groq: {'Ready ✅' if client else 'NOT SET ❌'}")
    print(f"  🧠  Model: {MODEL_NAME}")
    print("="*60+"\n")
    if not client:
        print("  1. Visit https://console.groq.com/keys")
        print("  2. Create free API key")
        print("  3. Set GROQ_API_KEY = 'your_key' at top of file\n")
    app.run(debug=True,host="0.0.0.0",port=5000)